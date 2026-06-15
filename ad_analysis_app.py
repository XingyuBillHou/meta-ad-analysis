# -*- coding: utf-8 -*-
"""
跨境电商广告投放多维度（日/周/月）智能分析与生成系统
作者: AI数据架构师
说明: 单文件 Streamlit 应用
"""

import re
import json
import os
import smtplib
import ssl
import base64
import tempfile
import time
from datetime import datetime, date, timedelta
from email.header import Header
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formataddr
from io import BytesIO
from pathlib import Path

import httpx
import numpy as np
import pandas as pd
import streamlit as st

# 避免 macOS 默认 locale 下 HTTP 请求体中文被按 ASCII 编码
os.environ.setdefault("PYTHONUTF8", "1")


# ============================================================
# 模块零：基础工具函数 - 日期识别与解析
# ============================================================

DATE_PATTERN = re.compile(r'^\d{4}[-/]\d{1,2}[-/]\d{1,2}')
WEEKEND_LABEL_RE = re.compile(r'周末|weekend', re.I)
SKIP_ROW_LABEL_RE = re.compile(
    r'^(week\s*\d*|february?|january|march|april|may|june|july|august|'
    r'september|october|november|december|日期|date)$',
    re.I,
)


def is_valid_date(val) -> bool:
    """
    判断一个单元格的值是否是"标准日期"。
    - pandas/python 的日期/时间类型 -> True
    - 形如 2026-06-09 / 2026/6/9 (可带时间后缀) 的字符串 -> True
    - 纯数字 (如周数 23) -> False
    - "Week" / "周末三日" 等纯文本 -> False
    """
    try:
        if val is None:
            return False
        if isinstance(val, float) and np.isnan(val):
            return False
        if pd.isna(val):
            return False
    except (TypeError, ValueError):
        pass

    if isinstance(val, (pd.Timestamp, datetime, date)):
        return True

    # 纯数字（int/float/numpy数字）不视为日期，避免把"周数"误判为日期
    if isinstance(val, (int, float, np.integer, np.floating)):
        return False

    s = str(val).strip()
    if not s:
        return False

    if DATE_PATTERN.match(s):
        try:
            pd.to_datetime(s)
            return True
        except Exception:
            return False
    return False


def to_date(val):
    """将一个值安全地转换为 datetime.date，转换失败返回 None"""
    try:
        if isinstance(val, (pd.Timestamp, datetime)):
            return val.date()
        if isinstance(val, date):
            return val
        return pd.to_datetime(str(val).strip()).date()
    except Exception:
        return None


def _md_dot_parts(val):
    """解析 M.DD / MM.DD 格式（如 11.13、9.1），返回 (month, day) 或 None。"""
    if isinstance(val, str):
        s = val.strip()
    elif isinstance(val, (int, float, np.integer, np.floating)) and not isinstance(val, bool):
        try:
            if isinstance(val, float) and np.isnan(val):
                return None
        except (TypeError, ValueError):
            pass
        s = f"{val:g}"
    else:
        return None

    if "." not in s:
        return None
    left, right = s.split(".", 1)
    try:
        month, day = int(left), int(right)
    except ValueError:
        return None
    if 1 <= month <= 12 and 1 <= day <= 31:
        return month, day
    return None


def is_md_dot_date(val) -> bool:
    """判断是否为 M.DD 短日期（Shopify / Nuage Wear 常用）。"""
    return _md_dot_parts(val) is not None


def _infer_year_for_month_day(month: int, day: int, ref_dates: list) -> int:
    if not ref_dates:
        return datetime.now().year
    years = sorted({d.year for d in ref_dates})
    best_year = years[-1]
    best_dist = None
    for year in years:
        try:
            target = date(year, month, day)
        except ValueError:
            continue
        dist = min(abs((target - rd).days) for rd in ref_dates)
        if best_dist is None or dist < best_dist:
            best_dist = dist
            best_year = year
    return best_year


def md_dot_to_date(val, ref_dates: list = None):
    parts = _md_dot_parts(val)
    if not parts:
        return None
    month, day = parts
    year = _infer_year_for_month_day(month, day, ref_dates or [])
    try:
        return date(year, month, day)
    except ValueError:
        return None


def parse_any_date(val, ref_dates: list = None):
    """统一解析标准日期、Timestamp、M.DD 短日期。"""
    if is_skip_row_label(val):
        return None
    if is_weekend_label(val):
        return None
    if is_valid_date(val):
        return to_date(val)
    if is_md_dot_date(val):
        return md_dot_to_date(val, ref_dates)
    return None


def is_skip_row_label(val) -> bool:
    """Week 汇总行、重复表头行等不参与日期匹配。"""
    try:
        if val is None or pd.isna(val):
            return True
    except (TypeError, ValueError):
        pass
    s = str(val).strip()
    if not s or s.lower() == "nan":
        return True
    return bool(SKIP_ROW_LABEL_RE.match(s))


def is_data_row(row) -> bool:
    """判断一行是否为数据行（首列为日期/周末汇总，且其余列有值）。"""
    if row is None or len(row) == 0:
        return False
    col0 = row.iloc[0]
    if is_weekend_label(col0):
        return True
    if is_skip_row_label(col0):
        return False
    if is_valid_date(col0) or is_md_dot_date(col0):
        return int(row.iloc[1:].notna().sum()) >= 1
    return False


def is_header_row(row) -> bool:
    """判断一行是否为表头行。"""
    if row is None or len(row) == 0:
        return False
    col0 = row.iloc[0]
    if pd.notna(col0) and str(col0).strip() in ("日期", "date", "Date", "DATE"):
        return True
    if is_data_row(row) or is_weekend_label(col0):
        return False
    text_count = sum(
        1 for v in row
        if pd.notna(v) and isinstance(v, str) and str(v).strip()
    )
    return text_count >= 2


def find_data_start_row(raw: pd.DataFrame) -> int:
    for i in range(len(raw)):
        if is_data_row(raw.iloc[i]):
            return i
    return 1


def find_header_start(raw: pd.DataFrame, data_start: int) -> int:
    header_start = data_start
    for i in range(data_start - 1, -1, -1):
        if is_header_row(raw.iloc[i]):
            header_start = i
        else:
            break
    return 0 if header_start == data_start and data_start > 0 else header_start


def collect_reference_dates(df: pd.DataFrame, date_col) -> list:
    refs = []
    for val in df[date_col]:
        if is_valid_date(val):
            d = to_date(val)
            if d:
                refs.append(d)
    return sorted(set(refs))


def build_parsed_date_series(df: pd.DataFrame, date_col) -> pd.Series:
    refs = collect_reference_dates(df, date_col)
    return df[date_col].apply(lambda x: parse_any_date(x, refs))


def is_weekend_label(val) -> bool:
    """判断是否为「周末三日」等汇总行标签（非标准日期）。"""
    try:
        if val is None:
            return False
        if isinstance(val, float) and np.isnan(val):
            return False
        if pd.isna(val):
            return False
    except (TypeError, ValueError):
        pass

    if is_valid_date(val):
        return False

    s = str(val).strip()
    if not s:
        return False
    return bool(WEEKEND_LABEL_RE.search(s)) or "三日" in s


def infer_weekend_range(cleaned_df: pd.DataFrame, date_col, row_idx: int):
    """
    根据「周末三日」行在表中的位置，推断对应的周五~周日日期。
    规则：通常紧跟在周四行之后；或以下一行周一反推。
    """
    for i in range(row_idx - 1, max(-1, row_idx - 8), -1):
        val = cleaned_df.iloc[i][date_col]
        refs = collect_reference_dates(cleaned_df, date_col)
        prev = parse_any_date(val, refs)
        if prev and prev.weekday() == 3:  # 周四
            return prev + timedelta(days=1), prev + timedelta(days=3)
        if prev:
            break

    for i in range(row_idx + 1, min(len(cleaned_df), row_idx + 8)):
        val = cleaned_df.iloc[i][date_col]
        refs = collect_reference_dates(cleaned_df, date_col)
        nxt = parse_any_date(val, refs)
        if nxt and nxt.weekday() == 0:  # 周一
            sun = nxt - timedelta(days=1)
            return sun - timedelta(days=2), sun
        if nxt:
            break

    return None, None


def get_weekend_buckets(sheets_dict: dict) -> list:
    """从 Shopify Sheet 扫描所有「周末三日」汇总段，返回可选 bucket 列表。"""
    try:
        shopify_df = get_shopify_df(sheets_dict)
        if shopify_df is None or shopify_df.empty:
            return []

        cleaned_df = clean_columns(shopify_df)
        date_col = find_date_column(cleaned_df)
        if date_col is None:
            return []

        buckets = []
        bucket_index = 0
        for idx in range(len(cleaned_df)):
            val = cleaned_df.iloc[idx][date_col]
            if not is_weekend_label(val):
                continue

            fri, sun = infer_weekend_range(cleaned_df, date_col, idx)
            if fri is None or sun is None:
                continue

            buckets.append({
                "bucket_index": bucket_index,
                "raw_label": str(val).strip(),
                "start": fri,
                "end": sun,
            })
            bucket_index += 1

        return buckets
    except Exception as e:
        st.warning(f"⚠️ 扫描周末三日数据时出错：{e}")
        return []


def _filter_weekend_rows(cleaned_df: pd.DataFrame, date_col, bucket: dict) -> pd.DataFrame:
    """在各 Sheet 中定位与 bucket 对应的「周末三日」汇总行。"""
    raw = cleaned_df[date_col]
    weekend_indices = [i for i, val in enumerate(raw) if is_weekend_label(val)]

    bucket_idx = bucket["bucket_index"]
    if bucket_idx < len(weekend_indices):
        return cleaned_df.iloc[[weekend_indices[bucket_idx]]]

    label = bucket["raw_label"]
    label_matches = cleaned_df[raw.astype(str).str.strip() == label]
    if bucket_idx < len(label_matches):
        return label_matches.iloc[[bucket_idx]]

    return pd.DataFrame()


# ============================================================
# 模块一(数据层)：复杂合并表头的"暴力拍平"读取引擎
# ============================================================

def flatten_merged_header(raw: pd.DataFrame) -> pd.DataFrame:
    """
    将已读出的原始表格（header=None）拍平 2-4 行合并表头，返回标准 DataFrame。
    适配 Shopify 等多层表头、M.DD 短日期等格式。
    """
    if raw.empty or raw.shape[0] < 1:
        return pd.DataFrame()

    data_start_row = find_data_start_row(raw)
    if data_start_row == 0:
        data_start_row = 1

    header_start = find_header_start(raw, data_start_row)
    header_block = raw.iloc[header_start:data_start_row]
    header_block = header_block.ffill(axis=1).ffill(axis=0)

    data_block = raw.iloc[data_start_row:].reset_index(drop=True)

    new_columns = []
    for col_idx in range(raw.shape[1]):
        parts = []
        for row_idx in range(len(header_block)):
            val = header_block.iloc[row_idx, col_idx]
            if pd.notna(val):
                s = str(val).strip()
                if s and s.lower() != "nan" and s not in parts:
                    parts.append(s)
        col_name = "_".join(parts) if parts else f"列{col_idx}"
        new_columns.append(col_name)

    seen = {}
    final_columns = []
    for c in new_columns:
        if c in seen:
            seen[c] += 1
            final_columns.append(f"{c}_{seen[c]}")
        else:
            seen[c] = 0
            final_columns.append(c)

    data_block.columns = final_columns
    return data_block


def smart_read_sheet(xls: pd.ExcelFile, sheet_name: str) -> pd.DataFrame:
    """通过 pandas ExcelFile 读取 Sheet 并拍平合并表头。"""
    raw = pd.read_excel(xls, sheet_name=sheet_name, header=None)
    return flatten_merged_header(raw)


def _read_uploaded_bytes(uploaded_file) -> bytes:
    uploaded_file.seek(0)
    return uploaded_file.read()


def _read_sheet_with_calamine(file_bytes: bytes, sheet_name: str) -> pd.DataFrame:
    """用 Rust 引擎 calamine 读取，可绕过 openpyxl 对损坏 sharedStrings 的报错。"""
    from python_calamine import CalamineWorkbook

    try:
        wb = CalamineWorkbook.from_filelike(BytesIO(file_bytes))
    except Exception:
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            wb = CalamineWorkbook.from_path(tmp_path)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    sheet = wb.get_sheet_by_name(sheet_name)
    if sheet is None:
        raise ValueError(f"Sheet「{sheet_name}」不存在")
    return pd.DataFrame(sheet.to_python())


def _open_excel_workbook(uploaded_file):
    """
    依次尝试 calamine / openpyxl 打开工作簿。
    返回 (sheet_names, read_raw_fn)，read_raw_fn(sheet_name) -> DataFrame(header=None)。
    """
    file_bytes = _read_uploaded_bytes(uploaded_file)
    errors = []

    for engine in ("calamine", "openpyxl"):
        try:
            xls = pd.ExcelFile(BytesIO(file_bytes), engine=engine)

            def read_raw(sheet_name, _xls=xls):
                return pd.read_excel(_xls, sheet_name=sheet_name, header=None)

            return xls.sheet_names, read_raw, engine
        except Exception as e:
            errors.append(f"{engine}: {e}")

    try:
        from python_calamine import CalamineWorkbook

        def read_raw(sheet_name, _bytes=file_bytes):
            return _read_sheet_with_calamine(_bytes, sheet_name)

        wb = CalamineWorkbook.from_filelike(BytesIO(file_bytes))
        return wb.sheet_names, read_raw, "calamine_direct"
    except Exception as e:
        errors.append(f"calamine_direct: {e}")

    detail = "\n".join(f"  - {msg}" for msg in errors)
    raise RuntimeError(
        "无法读取该 Excel 文件（常见于 WPS/在线表格导出的损坏 XML）。\n"
        f"{detail}\n"
        "建议：用 Microsoft Excel 打开后「另存为」新的 .xlsx 再上传。"
    )


# 未在投放的账号 Sheet，读取 Excel 时自动跳过
SKIPPED_SHEET_NAMES = frozenset({"nuage2026", "nuage wear"})


def _is_skipped_sheet(sheet_name: str) -> bool:
    return sheet_name.strip().lower() in SKIPPED_SHEET_NAMES


def load_all_sheets(uploaded_file) -> dict:
    """读取上传文件中的所有 Sheet（或 CSV 的单表），返回 {sheet名: DataFrame}"""
    sheets = {}
    try:
        file_name = uploaded_file.name.lower()

        if file_name.endswith(".csv"):
            try:
                df = pd.read_csv(uploaded_file)
            except UnicodeDecodeError:
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, encoding="gbk")
            sheets["数据表"] = df
            return sheets

        # xlsx 文件（多引擎容错：calamine -> openpyxl -> calamine 直读）
        sheet_names, read_raw, engine_used = _open_excel_workbook(uploaded_file)
        if engine_used != "openpyxl":
            st.info(f"ℹ️ 已使用 {engine_used} 引擎读取（原 openpyxl 无法解析该文件）。")

        skipped_sheets = []
        for sheet_name in sheet_names:
            if _is_skipped_sheet(sheet_name):
                skipped_sheets.append(sheet_name)
                continue
            try:
                raw = read_raw(sheet_name)
                df = flatten_merged_header(raw)
                if df is not None and not df.empty:
                    sheets[sheet_name] = df
                else:
                    st.warning(f"⚠️ Sheet「{sheet_name}」为空或无法解析，已跳过。")
            except Exception as e:
                st.warning(f"⚠️ 读取 Sheet「{sheet_name}」时发生错误，已跳过：{e}")

        if skipped_sheets:
            st.caption(
                "ℹ️ 已自动跳过未投放账号表："
                + "、".join(skipped_sheets)
            )

    except Exception as e:
        st.error(f"❌ 文件解析失败：{e}")

    return sheets


# ============================================================
# 模块二：纯净数据抓取引擎
# ============================================================

DROP_COL_KEYWORDS = [
    "总结", "分析", "理解", "复盘", "心得", "解读",
    "summary", "analysis", "insight", "comment", "note",
]


def clean_columns(df: pd.DataFrame) -> pd.DataFrame:
    """删除所有列名中包含'总结/分析/理解'等人工洞察字样的列"""
    if df is None or df.empty:
        return df

    cols_to_drop = []
    for col in df.columns:
        col_str = str(col).lower()
        for kw in DROP_COL_KEYWORDS:
            if kw.lower() in col_str:
                cols_to_drop.append(col)
                break

    if cols_to_drop:
        return df.drop(columns=cols_to_drop, errors="ignore")
    return df


def find_date_column(df: pd.DataFrame):
    """
    找到日期列：
    1. 优先匹配列名中含"日期"/"时间"/"date"的列；
    2. 否则在前几列中，选择"看起来像日期的值"最多的列；
    3. 实在找不到则返回第一列。
    """
    if df is None or df.empty or len(df.columns) == 0:
        return None

    keywords = ["日期", "时间", "date", "Date", "DATE"]
    for col in df.columns:
        col_str = str(col)
        for kw in keywords:
            if kw in col_str:
                return col

    best_col = df.columns[0]
    best_count = -1
    for col in df.columns[: min(3, len(df.columns))]:
        try:
            count = df[col].apply(is_valid_date).sum()
            if count > best_count:
                best_count = count
                best_col = col
        except Exception:
            continue

    return best_col


def extract_row_dict(row: pd.Series) -> dict:
    """把一行数据转成 {列名: 值} 的字典，剔除 NaN，并把 numpy/Timestamp 转成原生类型"""
    d = {}
    for col, val in row.items():
        try:
            if val is None:
                continue
            if isinstance(val, float) and np.isnan(val):
                continue
            if pd.isna(val):
                continue
        except (TypeError, ValueError):
            pass

        if isinstance(val, (pd.Timestamp, datetime, date)):
            val = str(val)[:10]
        elif isinstance(val, np.integer):
            val = int(val)
        elif isinstance(val, np.floating):
            fval = float(val)
            val = int(fval) if fval == int(fval) else round(fval, 4)
        elif isinstance(val, float):
            val = int(val) if val == int(val) else round(val, 4)

        d[str(col)] = val
    return d


def extract_data_for_report(
    sheets_dict: dict,
    date_mode: str,
    selected_date: date = None,
    start_date: date = None,
    end_date: date = None,
    weekend_bucket: dict = None,
) -> dict:
    """
    核心数据抓取：
    - single：各 Sheet 中日期 == selected_date 的行
    - range：日期落在 [start_date, end_date] 的行（时间序列）
    - weekend：匹配「周末三日」汇总行（周五~周日合并数据）
    """
    result = {}

    for sheet_name, df in sheets_dict.items():
        try:
            if df is None or df.empty:
                result[sheet_name] = f"【{sheet_name}】该 Sheet 数据为空，跳过。"
                continue

            cleaned_df = clean_columns(df.copy())
            date_col = find_date_column(cleaned_df)

            if date_col is None or date_col not in cleaned_df.columns:
                result[sheet_name] = f"【{sheet_name}】未识别到日期列，跳过该渠道数据提取。"
                continue

            if date_mode == "weekend":
                if not weekend_bucket:
                    result[sheet_name] = f"【{sheet_name}】未指定周末三日区间，跳过。"
                    continue
                filtered = _filter_weekend_rows(cleaned_df, date_col, weekend_bucket)
                # 同时附上该周末 Fri~Sun 的逐日明细（若有），供 AI 做趋势分析
                parsed_dates = pd.to_datetime(
                    build_parsed_date_series(cleaned_df, date_col)
                )
                fri, sun = weekend_bucket["start"], weekend_bucket["end"]
                daily_mask = (parsed_dates >= pd.Timestamp(fri)) & (
                    parsed_dates <= pd.Timestamp(sun)
                )
                daily_rows = cleaned_df[daily_mask]
                if not daily_rows.empty:
                    filtered = pd.concat([filtered, daily_rows]).drop_duplicates()
            else:
                parsed_dates = pd.to_datetime(
                    build_parsed_date_series(cleaned_df, date_col)
                )

                if date_mode == "single":
                    if selected_date is None:
                        result[sheet_name] = f"【{sheet_name}】未指定日期，跳过。"
                        continue
                    mask = parsed_dates == pd.Timestamp(selected_date)
                else:
                    if start_date is None or end_date is None:
                        result[sheet_name] = f"【{sheet_name}】未指定时间区间，跳过。"
                        continue
                    mask = (parsed_dates >= pd.Timestamp(start_date)) & (
                        parsed_dates <= pd.Timestamp(end_date)
                    )
                filtered = cleaned_df[mask]

            if filtered.empty:
                if date_mode == "weekend":
                    fri = weekend_bucket["start"].strftime("%Y-%m-%d")
                    sun = weekend_bucket["end"].strftime("%Y-%m-%d")
                    result[sheet_name] = (
                        f"【{sheet_name}】未找到 {fri} ~ {sun} 的「周末三日」汇总行。"
                    )
                else:
                    result[sheet_name] = f"【{sheet_name}】在此时间范围内无数据记录。"
                continue

            records = [extract_row_dict(row) for _, row in filtered.iterrows()]
            result[sheet_name] = records

        except Exception as e:
            result[sheet_name] = (
                f"处理 Sheet「{sheet_name}」时发生异常：{e}，已跳过该渠道数据，"
                f"不影响其他渠道分析。"
            )

    return result


# ============================================================
# 模块三：动态日期范围 - 基于 Shopify Sheet
# ============================================================

def get_shopify_df(sheets_dict: dict):
    """模糊匹配名字中含 shopify 的 Sheet；找不到则退而求其次用第一个 Sheet（兼容CSV单表场景）"""
    for name, df in sheets_dict.items():
        if "shopify" in str(name).lower():
            return df
    if sheets_dict:
        return list(sheets_dict.values())[0]
    return None


def get_valid_dates(sheets_dict: dict):
    """从 Shopify Sheet 的日期列中，提取所有符合 YYYY-MM-DD 的真实日期，排除纯文本汇总行"""
    try:
        shopify_df = get_shopify_df(sheets_dict)
        if shopify_df is None or shopify_df.empty:
            return []

        cleaned_df = clean_columns(shopify_df)
        date_col = find_date_column(cleaned_df)
        if date_col is None:
            return []

        valid_dates = set()
        parsed = build_parsed_date_series(cleaned_df, date_col)
        for d in parsed:
            if d:
                valid_dates.add(d)

        return sorted(valid_dates)
    except Exception as e:
        st.warning(f"⚠️ 提取日期列表时出错：{e}")
        return []


# ============================================================
# 模块四：AI 分析报告生成引擎
# ============================================================

def build_system_prompt(report_type: str, date_mode: str) -> str:
    if date_mode == "single":
        time_note = (
            f"这是一份【{report_type}】，数据为**单日**快照，"
            f"请重点关注该日各渠道的即时表现、异常波动及与近期均值的对比。"
        )
    elif date_mode == "weekend":
        time_note = (
            f"这是一份【{report_type}】，数据为**周末三日汇总**（周五+周六+周日合并），"
            f"请按整个周末区间评估表现，并与相邻工作日/周末对比。"
        )
    else:
        time_note = (
            f"这是一份【{report_type}】，数据为**日期范围内的时间序列**。"
            f"请务必从时间维度分析消耗、销售额、ROAS 等核心指标的变化趋势"
            f"（上升/下降/平稳/剧烈波动），并明确指出关键的转折点或异常日期。"
        )

    return f"""# 角色设定
你是一位拥有十年以上经验的顶尖跨境电商品牌 CMO（首席营销官），尤其精通 Meta、Google、AppLovin 等多渠道效果广告的数据分析与预算决策，风格数据驱动、直击要害、敢于给出明确结论和行动建议。

# 当前任务
请基于我提供的【原始投放数据】，生成一份【{report_type}】。
{time_note}

# 数据说明
- 数据来源可能包括：Shopify（电商大盘销售数据）、Google（Google Ads 投放数据）、Axon(AppLovin)（AppLovin 投放数据），以及多个 Meta 广告账号（如 WearNuage、Nuage Bra 等）。
- 每个渠道的数据是一份「记录列表」，每条记录对应一天（或一行）的原始字段，字段名以实际数据为准（可能包含消耗/Spend/Cost、曝光、点击、转化、ROAS、CPA、销售额/Revenue/Sales、订单数等）。
- 若某渠道在该区间无数据，会用文字注明"该渠道此区间无数据"，请正常处理，不要编造数据，也不要因此中断对其他渠道的分析。

# 输出结构要求（必须严格按以下 Markdown 结构输出，使用简体中文）

## 一、大盘纵览（Executive Summary）
- 结合 Shopify 的销售额（大盘数据）与各广告渠道的总消耗（前端总消耗），评估整体投入产出比（可估算"大盘 ROAS ≈ 总销售额 / 总广告消耗"）与业务体量。
- 若是周报/月报或日期范围/周末汇总，必须明确指出整体趋势方向（增长/下降/波动），并点出关键时间节点或异常日期。

## 二、渠道表现拆解
- 跨 Sheet 对比 Google、各个 Meta 账号、AppLovin（Axon）的消耗与 ROAS / 转化表现。
- 建议使用表格或分点方式，清晰指出：
  - 哪些渠道/账号是「增长引擎」（消耗增长且 ROAS 健康、转化稳健）；
  - 哪些渠道/账号在「拖后腿」（ROAS 偏低、消耗虚高或转化下滑）。

## 三、下一步行动指令（Action Items）
- 给出具体、可执行的建议，包括：
  - 预算调整：明确指出哪个渠道/账号 加预算多少（百分比或金额）、哪个 减预算多少；
  - 关停/暂停建议：对明显持续亏损、无改善趋势的渠道/账号给出关停建议；
  - 放量建议：对表现优异的渠道/账号给出加速放量的具体方式（如提升预算上限、扩量人群包等）。
- 建议必须具体可执行，避免"持续观察"、"进一步优化"等空泛表述。

# 注意事项
- 全文使用简体中文，语言专业、简洁、有决策力。
- 若某渠道数据缺失，简要说明即可，不影响其他部分的结论。
- 不要输出与上述结构无关的内容，不要重复粘贴原始数据。

# 硬性要求（必须遵守，否则视为失败）
- **必须完整输出三个章节**，不得在中途停止；第二章表格必须填写完整，不能只写表头。
- **所有结论必须引用原始数据中的具体数字**（消耗、ROAS、ROI、销售额、订单数等），禁止空泛描述。
- 第二章表格至少包含：Shopify 大盘、Google、Axon(AppLovin)、WearNuage、Nuage Bra 等所有有数据的渠道。
- 第三章至少给出 **3 条以上** 具体行动建议，每条含渠道名 + 量化调整幅度。
- 报告总篇幅不少于 **800 字**。
"""


def build_data_digest(extracted_data: dict) -> str:
    """生成关键指标 plaintext 摘要，确保模型能直接引用数字。"""
    lines = ["## 关键数据摘要（报告中必须引用以下数字）\n"]
    for sheet_name, data in extracted_data.items():
        if not isinstance(data, list) or not data:
            continue
        for idx, row in enumerate(data):
            label = f"记录{idx + 1}" if len(data) > 1 else "汇总"
            parts = []
            for k, v in row.items():
                if k == "日期" or v in (None, "", "/", "nan"):
                    continue
                try:
                    if isinstance(v, float) and np.isnan(v):
                        continue
                except (TypeError, ValueError):
                    pass
                parts.append(f"{k}={v}")
            if parts:
                lines.append(f"- **{sheet_name}** ({label}): " + "; ".join(parts))
    return "\n".join(lines) + "\n"


def _pick_summary_row(rows: list) -> dict:
    """优先取「周末三日」汇总行，否则取首行。"""
    for row in rows:
        if is_weekend_label(row.get("日期")):
            return row
    return rows[0] if rows else {}


def _is_numeric(val) -> bool:
    if val is None:
        return False
    if isinstance(val, (int, float, np.integer, np.floating)):
        try:
            return not (isinstance(val, float) and np.isnan(val))
        except (TypeError, ValueError):
            return True
    return False


def _fmt_metric(val) -> str:
    if not _is_numeric(val):
        return str(val)
    f = float(val)
    if f == int(f):
        return f"{int(f):,}"
    return f"{f:,.2f}"


def _find_first_numeric(row: dict, keywords: list, prefer: str = None):
    prefer_match = None
    first_match = None
    for key, val in row.items():
        if not _is_numeric(val):
            continue
        key_str = str(key)
        for kw in keywords:
            if kw.lower() in key_str.lower():
                if prefer and prefer in key_str:
                    prefer_match = float(val)
                elif first_match is None:
                    first_match = float(val)
                break
    return prefer_match if prefer_match is not None else first_match


def _get_shopify_row(extracted_data: dict) -> dict:
    for name, data in extracted_data.items():
        if "shopify" in str(name).lower() and isinstance(data, list) and data:
            return _pick_summary_row(data)
    return {}


def _shopify_spend_for_channel(shopify_row: dict, channel_name: str):
    if not shopify_row:
        return None
    name_lower = channel_name.lower().replace(" ", "")
    mapping = [
        ("google", "Ad Spent_Google"),
        ("axon", "Ad Spent_Axon(AppLovin)"),
        ("applovin", "Ad Spent_Axon(AppLovin)"),
        ("wearnuage", "Ad Spent_WearNuage"),
        ("nuagebra", "Ad Spent_NuageBra"),
        ("nuage2026", "Ad Spent_Nuage2026"),
        ("nuagewear", "Ad Spent_NuageWear"),
    ]
    for pattern, col in mapping:
        if pattern in name_lower:
            val = shopify_row.get(col)
            if _is_numeric(val):
                return float(val)
    return None


def _extract_channel_summary(sheet_name: str, row: dict, shopify_row: dict) -> dict:
    shopify_spend = _shopify_spend_for_channel(shopify_row, sheet_name)
    local_spend = _find_first_numeric(row, [
        "Ad Spent_Total", "总消耗", "Spent", "消耗", "Cost", "Spend",
    ])
    if shopify_spend is not None and "shopify" not in sheet_name.lower():
        spend = shopify_spend
    else:
        spend = local_spend

    roas = _find_first_numeric(row, ["ROAS", "ROI", "MER"], prefer="总体")
    if roas is None:
        roas = _find_first_numeric(row, ["总ROAS", "ROI_总"])
    orders = _find_first_numeric(row, ["出单量", "Orders", "订单", "Oders", "转化"])

    extras = []
    for key, val in row.items():
        if key == "日期" or not _is_numeric(val):
            continue
        kl = str(key).lower()
        if any(x in kl for x in ["消耗", "spent", "roas", "roi", "出单", "order", "订单"]):
            continue
        if ("cpm" in kl or "cpc" in kl) and "日期" not in str(key):
            extras.append(f"{key.split('_')[-1]}={_fmt_metric(val)}")

    return {
        "channel": sheet_name,
        "spend": spend,
        "roas": roas,
        "orders": orders,
        "extra": "; ".join(extras[:3]) if extras else "-",
    }


def build_channel_tables_markdown(extracted_data: dict) -> str:
    """用代码从原始数据生成完整渠道表格，不依赖 LLM。"""
    shopify_row = _get_shopify_row(extracted_data)
    lines = [
        "## 二、渠道表现拆解",
        "",
        "### 2.1 渠道核心指标汇总",
        "",
        "| 渠道/账号 | 消耗 (USD) | ROAS/ROI | 出单/订单 | 其他 |",
        "|----------|-----------|----------|----------|------|",
    ]

    summaries = []
    for sheet_name, data in extracted_data.items():
        if not isinstance(data, list) or not data:
            continue
        row = _pick_summary_row(data)
        if "shopify" in str(sheet_name).lower():
            s = _extract_channel_summary(sheet_name, row, shopify_row)
            s["spend"] = _find_first_numeric(row, ["Ad Spent_Total"])
            s["roas"] = _find_first_numeric(row, ["MER", "ROI", "ROAS"])
            s["orders"] = _find_first_numeric(row, ["订单", "Oders", "Orders"])
            sales = _find_first_numeric(row, ["Total Sales", "销售额"])
            s["extra"] = f"销售额={_fmt_metric(sales)}" if sales else "-"
            summaries.append(s)
        else:
            summaries.append(_extract_channel_summary(sheet_name, row, shopify_row))

    for s in summaries:
        spend_s = _fmt_metric(s["spend"]) if s["spend"] is not None else "-"
        roas_s = _fmt_metric(s["roas"]) if s["roas"] is not None else "-"
        orders_s = _fmt_metric(s["orders"]) if s["orders"] is not None else "-"
        lines.append(
            f"| {s['channel']} | {spend_s} | {roas_s} | {orders_s} | {s['extra']} |"
        )

    lines.extend(["", "### 2.2 各渠道完整指标明细", ""])
    for sheet_name, data in extracted_data.items():
        if not isinstance(data, list) or not data:
            lines.append(f"**{sheet_name}**：{data}")
            lines.append("")
            continue

        lines.append(f"#### {sheet_name}")
        lines.append("")
        lines.append("| 指标字段 | 数值 |")
        lines.append("|---------|------|")
        for row_idx, row in enumerate(data):
            prefix = ""
            if len(data) > 1:
                prefix = f"[{row.get('日期', f'记录{row_idx + 1}')}] "
            for key, val in row.items():
                if key == "日期":
                    continue
                if val in (None, "", "nan"):
                    continue
                try:
                    if isinstance(val, float) and np.isnan(val):
                        continue
                except (TypeError, ValueError):
                    pass
                display_val = "无数据" if val == "/" else val
                lines.append(f"| {prefix}{key} | {display_val} |")
        lines.append("")

    return "\n".join(lines)


def build_user_content(extracted_data: dict, date_info: str, date_mode: str) -> str:
    lines = [f"以下是【{date_info}】期间，各渠道/Sheet 的原始投放数据（已剔除人工总结/分析列）：\n"]
    lines.append(build_data_digest(extracted_data))

    for sheet_name, data in extracted_data.items():
        lines.append(f"\n## 渠道/Sheet：{sheet_name}\n")
        if isinstance(data, str):
            lines.append(data)
        else:
            try:
                json_str = json.dumps(data, ensure_ascii=False, indent=2, default=str)
                lines.append(f"```json\n{json_str}\n```")
            except Exception as e:
                lines.append(f"（该渠道数据序列化失败：{e}）")

    if date_mode == "range":
        lines.append(
            "\n\n请注意：以上每个渠道的数据均为按日期排列的列表（时间序列），"
            "请结合日期变化分析趋势。"
        )
    elif date_mode == "weekend":
        lines.append(
            "\n\n请注意：以上数据包含「周末三日」汇总行，以及该周末 Fri~Sun 的逐日明细（如有）。"
            "汇总行已含三天合并指标，请勿与逐日数据重复相加。"
        )

    return "\n".join(lines)


def _safe_report_filename(meta: dict, ext: str) -> str:
    date_part = meta.get("date_info", "").replace(" ", "_").replace("至", "to")
    return f"{meta.get('report_type', 'report')}_{date_part}.{ext}"


def build_full_markdown(report: str, meta: dict) -> str:
    """将 AI 输出整合为带封面信息的完整 Markdown 文档。"""
    generated_at = meta.get("generated_at") or datetime.now().strftime("%Y-%m-%d %H:%M")
    return f"""# 跨境电商广告投放分析报告

| 项目 | 内容 |
|------|------|
| 报告类型 | {meta.get('report_type', '')} |
| 数据区间 | {meta.get('date_info', '')} |
| 生成时间 | {generated_at} |

---

{report}
"""


def build_full_html(report: str, meta: dict) -> str:
    """将报告转为可打印的 HTML 文档。"""
    md_doc = build_full_markdown(report, meta)
    body_lines = []
    for line in md_doc.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            body_lines.append(f"<h1>{stripped[2:]}</h1>")
        elif stripped.startswith("## "):
            body_lines.append(f"<h2>{stripped[3:]}</h2>")
        elif stripped.startswith("### "):
            body_lines.append(f"<h3>{stripped[4:]}</h3>")
        elif stripped.startswith("|") and "---" not in stripped:
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            body_lines.append(
                "<tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>"
            )
        elif stripped == "---":
            body_lines.append("<hr>")
        elif stripped.startswith("- "):
            body_lines.append(f"<li>{stripped[2:]}</li>")
        elif stripped:
            body_lines.append(f"<p>{stripped}</p>")

    body_html = "\n".join(body_lines)
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{meta.get('report_type', '报告')} - {meta.get('date_info', '')}</title>
<style>
  body {{ font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif;
         max-width: 900px; margin: 40px auto; padding: 0 24px; line-height: 1.7; color: #222; }}
  h1 {{ border-bottom: 2px solid #1f77b4; padding-bottom: 8px; }}
  h2 {{ color: #1f77b4; margin-top: 28px; }}
  h3 {{ color: #444; }}
  table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
  td, th {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
  hr {{ border: none; border-top: 1px solid #eee; margin: 24px 0; }}
  li {{ margin: 4px 0; }}
</style>
</head>
<body>
{body_html}
</body>
</html>"""


def build_docx_bytes(report: str, meta: dict) -> bytes:
    """将报告转为 Word (.docx) 字节流。"""
    from docx import Document

    doc = Document()
    doc.add_heading("跨境电商广告投放分析报告", 0)
    doc.add_paragraph(f"报告类型：{meta.get('report_type', '')}")
    doc.add_paragraph(f"数据区间：{meta.get('date_info', '')}")
    doc.add_paragraph(
        f"生成时间：{meta.get('generated_at') or datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    doc.add_paragraph("")

    for line in report.split("\n"):
        text = line.rstrip()
        if not text:
            continue
        if text.startswith("## "):
            doc.add_heading(text[3:], level=1)
        elif text.startswith("### "):
            doc.add_heading(text[4:], level=2)
        elif text.startswith("- "):
            doc.add_paragraph(text[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", text):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", text), style="List Number")
        else:
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            doc.add_paragraph(clean)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$")

SMTP_PRESETS = {
    "Gmail": {"host": "smtp.gmail.com", "port": 587, "use_tls": True, "use_ssl": False},
    "QQ 邮箱": {"host": "smtp.qq.com", "port": 465, "use_tls": False, "use_ssl": True},
    "163 邮箱": {"host": "smtp.163.com", "port": 465, "use_tls": False, "use_ssl": True},
    "Outlook": {"host": "smtp.office365.com", "port": 587, "use_tls": True, "use_ssl": False},
    "自定义": {"host": "", "port": 587, "use_tls": True, "use_ssl": False},
}


def _sanitize_smtp_password(password: str) -> str:
    """去除授权码中误粘贴的空格（QQ/Gmail 授权码常见）。"""
    return re.sub(r"\s+", "", (password or "").strip())


def _smtp_auth_error_hint(cfg: dict, exc: Exception) -> str:
    host = cfg.get("host", "").lower()
    server_msg = ""
    if getattr(exc, "args", None):
        server_msg = " ".join(str(a) for a in exc.args if a)
    base = "❌ SMTP 登录失败。"
    if server_msg:
        base += f"\n\n服务器回复：`{server_msg}`"
    if "gmail.com" in host:
        base += (
            "\n\n**Gmail 请确认：**\n"
            "1. 已开启两步验证，并使用 **应用专用密码**（不是 Google 登录密码）\n"
            "2. 发件邮箱与登录账号一致\n"
            "3. 若在 Streamlit 云端部署，Gmail 可能拦截云服务器登录 — "
            "建议改用 **Resend API** 或 **QQ/163 邮箱**"
        )
    elif "qq.com" in host:
        base += (
            "\n\n**QQ 邮箱请确认：**\n"
            "1. 已在 QQ 邮箱设置 → 账户 中开启 SMTP 服务\n"
            "2. 使用的是 **16 位授权码**，不是 QQ 密码\n"
            "3. 发件邮箱填写完整地址，如 `name@qq.com`"
        )
    elif "163.com" in host:
        base += (
            "\n\n**163 邮箱请确认：**\n"
            "1. 已在邮箱设置中开启 SMTP / POP3\n"
            "2. 使用的是 **客户端授权码**，不是登录密码"
        )
    else:
        base += "\n\n请检查发件邮箱、授权码/密码是否正确。"
    return base


def _is_valid_email(addr: str) -> bool:
    return bool(EMAIL_RE.match((addr or "").strip()))


def _parse_email_list(text: str) -> list:
    addrs = []
    for part in re.split(r"[,;\s\n]+", text or ""):
        addr = part.strip()
        if addr and _is_valid_email(addr):
            addrs.append(addr)
    return addrs


def _get_smtp_config(override: dict = None) -> dict:
    """读取 SMTP 配置：优先使用界面传入，其次 Streamlit Secrets。"""
    override = override or {}
    port_raw = override.get("port") or _secret("email", "smtp_port") or "587"
    try:
        port = int(port_raw)
    except (TypeError, ValueError):
        port = 587

    user = (override.get("user") or _secret("email", "smtp_user")).strip()
    password = _sanitize_smtp_password(override.get("password") or _secret("email", "smtp_password"))
    host = (override.get("host") or _secret("email", "smtp_host")).strip()
    from_addr = (override.get("from_addr") or _secret("email", "from_addr") or user).strip()

    use_ssl = override.get("use_ssl")
    if use_ssl is None:
        use_ssl = _secret("email", "use_ssl", default="").lower() == "true"
    use_tls = override.get("use_tls")
    if use_tls is None:
        use_tls = _secret("email", "use_tls", default="true").lower() != "false"

    if port == 465 and not use_ssl:
        use_ssl = True
        use_tls = False

    if not host or not user or not password or not from_addr:
        return {}
    return {
        "host": host,
        "port": port,
        "user": user,
        "password": password,
        "from_addr": from_addr,
        "from_name": override.get("from_name") or _secret("email", "from_name") or "跨境电商广告分析报告",
        "use_tls": use_tls and not use_ssl,
        "use_ssl": use_ssl,
    }


def _smtp_connect_and_send(host: str, port: int, use_ssl: bool, use_tls: bool,
                           user: str, password: str, msg: MIMEMultipart) -> None:
    context = ssl.create_default_context()
    if use_ssl:
        with smtplib.SMTP_SSL(host, port, timeout=30, context=context) as smtp:
            smtp.ehlo()
            smtp.login(user, password)
            smtp.send_message(msg)
        return
    with smtplib.SMTP(host, port, timeout=30) as smtp:
        smtp.ehlo()
        if use_tls:
            smtp.starttls(context=context)
            smtp.ehlo()
        smtp.login(user, password)
        smtp.send_message(msg)


def _smtp_login_and_send(cfg: dict, msg: MIMEMultipart) -> None:
    user = cfg["user"].strip()
    password = _sanitize_smtp_password(cfg["password"])
    attempts = [(cfg["host"], cfg["port"], cfg["use_ssl"], cfg["use_tls"])]

    host_lower = cfg["host"].lower()
    if "gmail.com" in host_lower and not cfg["use_ssl"]:
        attempts.append((cfg["host"], 465, True, False))
    if "qq.com" in host_lower and cfg["use_ssl"]:
        attempts.append((cfg["host"], 587, False, True))

    last_auth_err = None
    for host, port, use_ssl, use_tls in attempts:
        try:
            _smtp_connect_and_send(host, port, use_ssl, use_tls, user, password, msg)
            return
        except smtplib.SMTPAuthenticationError as e:
            last_auth_err = e
            continue
    if last_auth_err:
        raise last_auth_err
    _smtp_connect_and_send(
        cfg["host"], cfg["port"], cfg["use_ssl"], cfg["use_tls"], user, password, msg
    )


def _send_via_resend(
    to_addrs: list,
    report: str,
    meta: dict,
    attachment_format: str,
    resend_api_key: str,
    from_addr: str,
    from_name: str,
) -> tuple:
    if not resend_api_key.strip():
        return False, "请填写 Resend API Key（在 resend.com 注册后获取）。"
    if not from_addr.strip():
        return False, "请填写发件地址（Resend 需已验证的域名邮箱，测试可用 onboarding@resend.dev）。"

    filename, payload, mime_type = _build_email_attachment(report, meta, attachment_format)
    subject = f"{meta.get('report_type', '分析报告')} - {meta.get('date_info', '')}"
    generated_at = meta.get("generated_at") or datetime.now().strftime("%Y-%m-%d %H:%M")
    html_body = (
        f"<p>您好，</p>"
        f"<p>附件为自动生成的跨境电商广告投放分析报告。</p>"
        f"<ul>"
        f"<li>报告类型：{meta.get('report_type', '')}</li>"
        f"<li>数据区间：{meta.get('date_info', '')}</li>"
        f"<li>生成时间：{generated_at}</li>"
        f"</ul>"
        f"<p>完整内容请查看邮件附件。</p>"
    )

    body = {
        "from": f"{from_name} <{from_addr.strip()}>",
        "to": to_addrs,
        "subject": subject,
        "html": html_body,
        "attachments": [{
            "filename": filename,
            "content": base64.b64encode(payload).decode("ascii"),
        }],
    }

    try:
        resp = httpx.post(
            "https://api.resend.com/emails",
            headers={
                "Authorization": f"Bearer {resend_api_key.strip()}",
                "Content-Type": "application/json",
            },
            json=body,
            timeout=30.0,
        )
        if resp.status_code in (200, 201):
            return True, f"✅ 报告已通过 Resend 发送至：{', '.join(to_addrs)}"
        detail = resp.text[:500]
        return False, f"❌ Resend 发送失败 (HTTP {resp.status_code})：{detail}"
    except Exception as e:
        return False, f"❌ Resend 发送失败：{e}"


def _build_email_attachment(report: str, meta: dict, attachment_format: str):
    base_name = _safe_report_filename(meta, "").rstrip(".")
    if attachment_format == "html":
        return (
            f"{base_name}.html",
            build_full_html(report, meta).encode("utf-8"),
            "text/html",
        )
    if attachment_format == "md":
        return (
            f"{base_name}.md",
            build_full_markdown(report, meta).encode("utf-8"),
            "text/markdown",
        )
    return (
        f"{base_name}.docx",
        build_docx_bytes(report, meta),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def send_report_email(
    to_addrs: list,
    report: str,
    meta: dict,
    attachment_format: str = "docx",
    smtp_override: dict = None,
    email_method: str = "smtp",
    resend_api_key: str = "",
    resend_from_addr: str = "",
    resend_from_name: str = "跨境电商广告分析报告",
) -> tuple:
    """发送报告附件。返回 (成功与否, 提示信息)。"""
    if email_method == "resend":
        return _send_via_resend(
            to_addrs, report, meta, attachment_format,
            resend_api_key, resend_from_addr, resend_from_name,
        )

    cfg = _get_smtp_config(smtp_override)
    if not cfg:
        return False, (
            "未配置邮件服务器。请在左侧侧边栏「邮件发信配置」中填写 SMTP 信息，"
            "或在 Streamlit Secrets 中设置 email.smtp_host / smtp_user / smtp_password。"
        )
    if not to_addrs:
        return False, "请填写至少一个有效的收件人邮箱。"

    filename, payload, mime_type = _build_email_attachment(report, meta, attachment_format)
    subject = f"{meta.get('report_type', '分析报告')} - {meta.get('date_info', '')}"
    generated_at = meta.get("generated_at") or datetime.now().strftime("%Y-%m-%d %H:%M")

    msg = MIMEMultipart()
    msg["Subject"] = Header(subject, "utf-8")
    msg["From"] = formataddr((cfg["from_name"], cfg["from_addr"]))
    msg["To"] = ", ".join(to_addrs)

    body = f"""您好，

附件为自动生成的跨境电商广告投放分析报告。

报告类型：{meta.get('report_type', '')}
数据区间：{meta.get('date_info', '')}
生成时间：{generated_at}

此邮件由分析系统自动发送，请勿直接回复。
"""
    msg.attach(MIMEText(body, "plain", "utf-8"))

    attachment = MIMEApplication(payload, Name=filename)
    attachment.add_header("Content-Disposition", "attachment", filename=("utf-8", "", filename))
    attachment.add_header("Content-Type", mime_type)
    msg.attach(attachment)

    try:
        _smtp_login_and_send(cfg, msg)
        return True, f"✅ 报告已发送至：{', '.join(to_addrs)}"
    except smtplib.SMTPAuthenticationError as e:
        return False, _smtp_auth_error_hint(cfg, e)
    except smtplib.SMTPException as e:
        return False, f"❌ 邮件发送失败（SMTP）：{e}"
    except OSError as e:
        return False, f"❌ 无法连接邮件服务器 {cfg['host']}:{cfg['port']}：{e}"
    except Exception as e:
        return False, f"❌ 邮件发送失败：{e}"


def _require_ascii(value: str, field_name: str) -> str:
    """API Key / Base URL / 模型名必须为 ASCII，防止误粘贴中文说明。"""
    value = (value or "").strip()
    if not value:
        return value
    try:
        value.encode("ascii")
    except UnicodeEncodeError:
        raise ValueError(
            f"「{field_name}」包含非英文字符，请只填写 Key / URL / 模型名，"
            f"不要粘贴中文说明文字。"
        )
    return value


def _validate_llm_credentials(api_key: str, base_url: str) -> None:
    """在发起请求前检测常见 Key / URL 错配。"""
    key = api_key.strip().lower()

    if key.startswith("apify_api") or key.startswith("apify_ap"):
        raise ValueError(
            "检测到 **Apify API Key**，不能用于 AI 报告生成。"
            "请检查 Gemini API Key 配置。"
        )


def _sanitize_api_key(api_key: str) -> str:
    """去除首尾空格及误粘贴的引号。"""
    return api_key.strip().strip('"').strip("'").strip()


def _format_api_error(status: int, detail: str, api_key: str, base_url: str) -> str:
    """将 HTTP 错误转为更易读的中文提示。"""
    detail_lower = detail.lower()
    base_lower = base_url.strip().lower()
    key_hint = ""

    if "apify" in detail_lower or api_key.strip().lower().startswith("apify"):
        key_hint = (
            "\n\n**原因**：API Key 配置有误，请检查 Gemini Key。"
        )
    elif status == 429:
        if "generativelanguage.googleapis.com" in base_lower or "quota" in detail_lower:
            key_hint = (
                "\n\n**原因**：Google Gemini 免费额度已用完或未开通（limit: 0 表示当前无可用配额）。\n\n"
                "**可选方案**：\n"
                "1. 在 [Google AI Studio](https://aistudio.google.com/) 关联账单\n"
                "2. 将模型改为 `gemini-2.5-flash-lite` 后重试\n"
                "3. 查看用量：[ai.dev/rate-limit](https://ai.dev/rate-limit)"
            )
        else:
            key_hint = "\n\n**原因**：请求频率或配额超限，请稍后重试或更换模型。"
    elif status == 401:
        key_hint = "\n\n**原因**：Gemini API Key 无效，请核对侧边栏配置。"
    elif status == 404:
        key_hint = "\n\n**原因**：模型名称可能不正确。"
    elif status in (502, 503, 529):
        key_hint = (
            "\n\n**原因**：Gemini 服务暂时不可用（高峰期常见，通常几分钟内恢复）。\n\n"
            "**建议**：\n"
            "1. 等待 1~2 分钟后点击「生成 AI 分析报告」重试\n"
            "2. 改用 `gemini-2.5-flash-lite` 等轻量模型"
        )

    return f"❌ API 请求失败 (HTTP {status})：{detail}{key_hint}"


def _request_chat_completion(
    client,
    url: str,
    headers: dict,
    payload: dict,
    retry_delays: tuple,
) -> dict:
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    for attempt in range(len(retry_delays) + 1):
        try:
            resp = client.post(url, content=body, headers=headers)
            if resp.status_code in (429, 502, 503, 529) and attempt < len(retry_delays):
                time.sleep(retry_delays[attempt])
                continue
            resp.raise_for_status()
            return resp.json()
        except httpx.TimeoutException:
            if attempt < len(retry_delays):
                time.sleep(retry_delays[attempt])
                continue
            raise
    raise RuntimeError("API 多次重试后仍失败")


REPORT_SECTIONS = [
    {
        "title": "## 一、大盘纵览（Executive Summary）",
        "instruction": (
            "请**只写第一章**「大盘纵览（Executive Summary）」。\n"
            "- 引用 Shopify 总销售额、总广告消耗、订单数、MER/ROI 等**具体数字**\n"
            "- 计算并写出大盘 ROAS（销售额÷广告消耗）\n"
            "- 3~5 条要点，评估本周期整体盈亏与体量\n"
            "**禁止**写第二、三章。"
        ),
        "min_chars": 180,
    },
    {
        "id": "channel",
        "title": "## 二、渠道表现拆解",
        "instruction": (
            "渠道数据表格已由系统自动生成（见下方），你**不要重复输出表格**。\n"
            "请**只写**「### 2.3 渠道洞察与诊断」小节，包含：\n"
            "- 3~5 条 bullet：**增长引擎**（渠道名 + 具体 ROAS/消耗数字 + 原因）\n"
            "- 3~5 条 bullet：**拖后腿渠道**（渠道名 + 具体问题数字 + 风险）\n"
            "- 2~3 条 cross-channel 对比结论（如 Meta vs Google vs AppLovin）\n"
            "**禁止**写第一、三章；**禁止**重新输出 Markdown 表格。"
        ),
        "min_chars": 200,
        "skip_title": True,
    },
    {
        "title": "## 三、下一步行动指令（Action Items）",
        "instruction": (
            "请**只写第三章**「下一步行动指令（Action Items）」。\n"
            "- 至少 **5 条** numbered list，每条格式：\n"
            "  **【渠道/账号】** 动作描述（含量化幅度，如 +20% / -500 USD / 暂停 XX 广告组）\n"
            "- 覆盖：加预算、减预算、暂停、放量 等类型\n"
            "**禁止**写第一、二章。"
        ),
        "min_chars": 280,
    },
]


def _normalize_section_output(title: str, content: str) -> str:
    text = (content or "").strip()
    if not text:
        return f"{title}\n\n（该章节生成失败，请重试。）"
    if not text.startswith("##"):
        text = f"{title}\n\n{text}"
    elif title.split("（")[0] not in text.split("\n")[0]:
        text = f"{title}\n\n{text}"
    return text


def generate_report(
    api_key: str,
    base_url: str,
    model_name: str,
    system_prompt: str,
    user_content: str,
    temperature: float = 0.5,
    extracted_data: dict = None,
) -> str:
    """
    分三节依次调用大模型，再合并为完整报告，避免单次输出被截断。
    """
    read_timeout = 300.0
    connect_timeout = 60.0
    retry_delays = (5, 15, 30)
    max_output_tokens = 4096

    if not api_key:
        return "❌ 错误：请先在左侧侧边栏填写 Gemini API Key。"

    api_key = _require_ascii(_sanitize_api_key(api_key), "API Key")
    base_url = _require_ascii(base_url.rstrip("/"), "Base URL")
    model_name = _require_ascii(model_name, "模型名称")
    _validate_llm_credentials(api_key, base_url)

    url = f"{base_url}/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json; charset=utf-8",
    }
    timeout = httpx.Timeout(read_timeout, connect=connect_timeout)
    report_temperature = min(temperature, 0.4)

    sections_out = []
    try:
        with httpx.Client(timeout=timeout) as client:
            for section in REPORT_SECTIONS:
                if section.get("id") == "channel" and extracted_data:
                    channel_tables = build_channel_tables_markdown(extracted_data)
                    section_user = (
                        f"{user_content}\n\n---\n已生成的渠道表格：\n{channel_tables}\n\n"
                        f"{section['instruction']}"
                    )
                else:
                    section_user = f"{user_content}\n\n---\n{section['instruction']}"

                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": section_user},
                ]

                content = ""
                for attempt in range(2):
                    payload = {
                        "model": model_name,
                        "messages": messages,
                        "temperature": report_temperature,
                        "max_tokens": max_output_tokens,
                    }
                    data = _request_chat_completion(
                        client, url, headers, payload, retry_delays
                    )
                    choice = data.get("choices", [{}])[0]
                    content = choice.get("message", {}).get("content") or ""
                    finish_reason = choice.get("finish_reason", "")

                    if (
                        content
                        and len(content.strip()) >= section["min_chars"]
                        and finish_reason != "length"
                    ):
                        break

                    if attempt == 0:
                        messages.append({
                            "role": "assistant",
                            "content": content or "（空）",
                        })
                        messages.append({
                            "role": "user",
                            "content": (
                                "内容太短或不完整。请重新撰写本节，"
                                "必须引用上方数据中的**全部关键数字**，写得更详细、更完整。"
                            ),
                        })

                if section.get("id") == "channel" and extracted_data:
                    analysis = content.strip()
                    if analysis.startswith("##"):
                        analysis = analysis.split("\n", 1)[-1].strip()
                    sections_out.append(
                        channel_tables + "\n\n" + (
                            analysis if analysis.startswith("###")
                            else "### 2.3 渠道洞察与诊断\n\n" + analysis
                        )
                    )
                else:
                    sections_out.append(
                        _normalize_section_output(section["title"], content)
                    )

        return "\n\n".join(sections_out)

    except ValueError as e:
        return f"❌ {e}"
    except httpx.TimeoutException:
        partial = "\n\n".join(sections_out)
        if partial:
            return partial + (
                "\n\n---\n⚠️ *后续章节因超时未完成，请缩小日期范围或换用 deepseek-chat 后重试。*"
            )
        return (
            "❌ 模型响应超时。\n\n"
            "**建议**：换用 deepseek-chat / gpt-4o-mini，或缩小数据范围后重试。"
        )
    except httpx.HTTPStatusError as e:
        detail = e.response.text[:500] if e.response is not None else str(e)
        status = e.response.status_code if e.response is not None else 0
        return _format_api_error(status, detail, api_key, base_url)
    except Exception as e:
        return f"❌ 调用大模型 API 时发生错误：{e}\n\n请检查 API Key、Base URL、模型名称是否正确。"


def call_llm(api_key: str, base_url: str, model_name: str,
              system_prompt: str, user_content: str, temperature: float = 0.5,
              extracted_data: dict = None) -> str:
    """兼容入口：委托给分节生成。"""
    return generate_report(
        api_key, base_url, model_name, system_prompt, user_content, temperature,
        extracted_data=extracted_data,
    )


# ============================================================
# Streamlit Cloud Secrets（可选）
# ============================================================

def _secret(*keys, default=""):
    """从 Streamlit Cloud「Settings → Secrets」读取配置，本地无 secrets 时返回 default。"""
    try:
        node = st.secrets
        for key in keys:
            node = node[key]
        if node is None:
            return default
        return str(node).strip()
    except (KeyError, TypeError, AttributeError, FileNotFoundError):
        return default


# ============================================================
# Google Gemini 配置（唯一大模型）
# ============================================================

GEMINI_BASE_URL = "https://generativelanguage.googleapis.com/v1beta/openai"
GEMINI_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-flash-lite",
    "gemini-2.5-pro",
    "gemini-2.0-flash",
]
GEMINI_DEFAULT_MODEL = GEMINI_MODELS[0]


PROVIDER_MODELS = {
    "Google Gemini": GEMINI_MODELS,
}


# ============================================================
# 主程序：Streamlit UI
# ============================================================

def main():
    st.set_page_config(
        page_title="跨境电商广告智能分析系统",
        page_icon="📊",
        layout="wide",
    )

    st.title("📊 跨境电商广告投放多维度智能分析与生成系统")
    st.caption("支持日报 / 周报 / 月报，自动跨 Shopify、Google、Meta、AppLovin 多渠道数据生成 AI 分析报告")

    # ---------------- 侧边栏：Gemini 配置 ----------------
    with st.sidebar:
        st.header("⚙️ Gemini 配置")

        api_key = st.text_input(
            "API Key",
            type="password",
            placeholder="Gemini API Key",
            help="从 https://aistudio.google.com/apikey 获取",
        )
        base_url = GEMINI_BASE_URL

        default_model = GEMINI_DEFAULT_MODEL
        model_name = st.selectbox(
            "模型",
            options=GEMINI_MODELS,
            index=GEMINI_MODELS.index(default_model),
            help="推荐 gemini-2.5-flash；配额紧张时可改用 flash-lite。",
        )

        st.caption(
            "💡 Gemini 免费额度有限；若报 429，请在 "
            "[AI Studio](https://aistudio.google.com/) 激活账单，或改用 flash-lite 模型。"
        )

        temperature = st.slider("Temperature（创意度）", 0.0, 1.0, 0.5, 0.1)

        st.markdown("---")
        with st.expander("📧 邮件发信配置", expanded=False):
            email_method = st.radio(
                "发信方式",
                options=["smtp", "resend"],
                format_func=lambda x: "SMTP 邮箱" if x == "smtp" else "Resend API（推荐 Streamlit 云端）",
                help="Streamlit 云端部署时，Gmail SMTP 常被拦截，建议用 Resend。",
            )

            if email_method == "resend":
                resend_api_key = st.text_input(
                    "Resend API Key",
                    type="password",
                    key="resend_api_key",
                    placeholder="re_xxxxxxxx",
                    help="在 https://resend.com 注册获取，免费 100 封/天。",
                )
                resend_from = st.text_input(
                    "发件地址",
                    key="resend_from_addr",
                    value=_secret("email", "resend_from") or "onboarding@resend.dev",
                    help="测试可用 onboarding@resend.dev；正式使用需验证自己的域名。",
                )
                smtp_preset = "自定义"
                smtp_host = smtp_port = smtp_user = smtp_password = smtp_from = ""
                use_ssl = use_tls = False
                st.caption("Resend 通过 HTTPS 发信，不依赖 SMTP，适合 Streamlit Cloud。")
            else:
                resend_api_key = ""
                resend_from = ""
                smtp_preset = st.selectbox(
                    "邮箱类型",
                    options=list(SMTP_PRESETS.keys()),
                    help="Gmail/QQ/163 需使用授权码，不是登录密码。",
                )
                preset = SMTP_PRESETS[smtp_preset]
                if smtp_preset == "自定义":
                    smtp_host = st.text_input(
                        "SMTP 服务器",
                        key="smtp_host",
                        placeholder="smtp.example.com",
                    )
                    smtp_port = st.number_input("端口", min_value=1, max_value=65535, value=587)
                    use_ssl = st.checkbox("使用 SSL（465 端口）", value=False)
                    use_tls = st.checkbox("使用 STARTTLS（587 端口）", value=not use_ssl)
                else:
                    smtp_host = preset["host"]
                    smtp_port = preset["port"]
                    use_ssl = preset["use_ssl"]
                    use_tls = preset["use_tls"]
                    st.caption(f"SMTP：`{smtp_host}`　端口：`{smtp_port}`")

                smtp_user = st.text_input(
                    "发件邮箱",
                    key="smtp_user",
                    placeholder="your@gmail.com",
                )
                smtp_password = st.text_input(
                    "授权码 / 密码",
                    type="password",
                    key="smtp_password",
                    help="Gmail：应用专用密码；QQ/163：SMTP 授权码（复制后如有空格会自动去除）。",
                )
                smtp_from = st.text_input(
                    "发件地址（通常与发件邮箱相同）",
                    key="smtp_from_addr",
                    placeholder="与发件邮箱相同",
                )

        smtp_override = {
            "host": smtp_host if email_method == "smtp" and smtp_preset == "自定义" else (
                SMTP_PRESETS.get(smtp_preset, {}).get("host", "") if email_method == "smtp" else ""
            ),
            "port": int(smtp_port) if email_method == "smtp" else 587,
            "user": smtp_user if email_method == "smtp" else "",
            "password": smtp_password if email_method == "smtp" else "",
            "from_addr": (smtp_from or smtp_user) if email_method == "smtp" else "",
            "use_tls": use_tls if email_method == "smtp" and smtp_preset == "自定义" else (
                SMTP_PRESETS.get(smtp_preset, {}).get("use_tls", True) if email_method == "smtp" else True
            ),
            "use_ssl": use_ssl if email_method == "smtp" and smtp_preset == "自定义" else (
                SMTP_PRESETS.get(smtp_preset, {}).get("use_ssl", False) if email_method == "smtp" else False
            ),
        }
        if email_method == "resend":
            smtp_ready = bool(resend_api_key.strip() and resend_from.strip())
        else:
            smtp_ready = bool(_get_smtp_config(smtp_override))

    # ---------------- 主区：文件上传 ----------------
    uploaded_file = st.file_uploader("📁 上传广告投放数据文件 (.xlsx / .csv)", type=["xlsx", "csv"])

    if uploaded_file is None:
        st.info("请上传包含 Shopify / Google / Meta / AppLovin 等 Sheet 的 Excel 文件，或单表 CSV 文件。")
        return

    # 文件变更时重置缓存
    if (
        "file_id" not in st.session_state
        or st.session_state["file_id"] != (uploaded_file.name, uploaded_file.size)
    ):
        with st.spinner("正在解析文件，自动处理合并表头..."):
            sheets = load_all_sheets(uploaded_file)
        st.session_state["sheets"] = sheets
        st.session_state["file_id"] = (uploaded_file.name, uploaded_file.size)
        st.session_state.pop("report", None)  # 新文件，清空旧报告

    sheets = st.session_state.get("sheets", {})

    if not sheets:
        st.error("❌ 未能从文件中解析出任何有效数据，请检查文件格式。")
        return

    st.success(f"✅ 成功解析 {len(sheets)} 个数据表：{', '.join(sheets.keys())}")

    with st.expander("🔍 查看各 Sheet 数据预览（前5行，已剔除总结/分析列）"):
        for name, df in sheets.items():
            st.markdown(f"**{name}**  （共 {len(df)} 行，{len(df.columns)} 列）")
            try:
                st.dataframe(clean_columns(df).head(5), use_container_width=True)
            except Exception as e:
                st.warning(f"预览失败：{e}")

    # ---------------- 报告维度 + 动态日期选择 ----------------
    st.markdown("---")
    st.subheader("📅 报告参数设置")

    col_a, col_b = st.columns([1, 2])
    with col_a:
        report_type = st.selectbox("选择报告维度", ["日报", "周报", "月报"])

    valid_dates = get_valid_dates(sheets)
    weekend_buckets = get_weekend_buckets(sheets)

    if not valid_dates and not weekend_buckets:
        st.error("❌ 无法从数据中提取到有效日期或周末汇总行，请检查 Shopify Sheet。")
        return

    date_mode_options = ["单一日期", "日期范围"]
    if weekend_buckets:
        date_mode_options.append("周末三日")

    selected_date = None
    start_date = None
    end_date = None
    weekend_bucket = None
    date_info = ""
    date_mode = "single"

    with col_b:
        date_mode_label = st.radio(
            "数据范围",
            date_mode_options,
            horizontal=True,
            help="单一日期：选一天；日期范围：自定义起止；周末三日：选周五~周日汇总行",
        )
        date_mode_map = {"单一日期": "single", "日期范围": "range", "周末三日": "weekend"}
        date_mode = date_mode_map[date_mode_label]

        if date_mode == "single":
            if not valid_dates:
                st.warning("⚠️ 未找到标准日期，请改用「日期范围」或「周末三日」。")
            else:
                selected_date = st.selectbox(
                    "选择日期",
                    options=valid_dates,
                    index=len(valid_dates) - 1,
                    format_func=lambda d: d.strftime("%Y-%m-%d (%a)"),
                )
                start_date = end_date = selected_date
                date_info = selected_date.strftime("%Y-%m-%d")

        elif date_mode == "range":
            if not valid_dates:
                st.warning("⚠️ 未找到标准日期，无法设定范围。")
            else:
                min_d, max_d = min(valid_dates), max(valid_dates)
                c1, c2 = st.columns(2)
                with c1:
                    start_date = st.date_input(
                        "开始日期", value=min_d, min_value=min_d, max_value=max_d
                    )
                with c2:
                    end_date = st.date_input(
                        "结束日期", value=max_d, min_value=min_d, max_value=max_d
                    )
                if start_date > end_date:
                    st.warning("⚠️ 开始日期晚于结束日期，已自动交换。")
                    start_date, end_date = end_date, start_date
                if start_date == end_date:
                    date_info = start_date.strftime("%Y-%m-%d")
                else:
                    date_info = f"{start_date.strftime('%Y-%m-%d')} 至 {end_date.strftime('%Y-%m-%d')}"

        else:  # weekend
            def _format_weekend_bucket(b):
                return (
                    f"周末三日 ({b['start'].strftime('%Y-%m-%d')} ~ "
                    f"{b['end'].strftime('%Y-%m-%d')})"
                )

            weekend_bucket = st.selectbox(
                "选择周末三日",
                options=weekend_buckets,
                index=len(weekend_buckets) - 1,
                format_func=_format_weekend_bucket,
            )
            start_date = weekend_bucket["start"]
            end_date = weekend_bucket["end"]
            date_info = _format_weekend_bucket(weekend_bucket)

    if valid_dates:
        st.caption(
            f"📌 标准日期：{min(valid_dates)} ~ {max(valid_dates)}，共 {len(valid_dates)} 天"
        )
    if weekend_buckets:
        st.caption(f"📌 检测到 {len(weekend_buckets)} 个「周末三日」汇总段")

    # ---------------- 生成报告 ----------------
    st.markdown("---")
    generate_btn = st.button("🚀 生成 AI 分析报告", type="primary", use_container_width=False)

    if generate_btn:
        if not api_key:
            st.error("❌ 请先在左侧侧边栏填写 Gemini API Key。")
        else:
            try:
                with st.spinner("正在提取多渠道数据..."):
                    extracted_data = extract_data_for_report(
                        sheets_dict=sheets,
                        date_mode=date_mode,
                        selected_date=selected_date,
                        start_date=start_date,
                        end_date=end_date,
                        weekend_bucket=weekend_bucket,
                    )

                with st.expander("🧾 查看提交给 AI 的原始数据（调试用）"):
                    for sheet_name, data in extracted_data.items():
                        st.markdown(f"**{sheet_name}**")
                        if isinstance(data, str):
                            st.text(data)
                        else:
                            try:
                                st.code(
                                    json.dumps(data, ensure_ascii=False, indent=2, default=str),
                                    language="json",
                                )
                            except Exception as e:
                                st.text(f"（展示失败：{e}）")

                with st.spinner("AI 正在分三节生成完整报告（约 2~5 分钟）..."):
                    system_prompt = build_system_prompt(report_type, date_mode)
                    user_content = build_user_content(extracted_data, date_info, date_mode)
                    report = generate_report(
                        api_key=api_key,
                        base_url=base_url,
                        model_name=model_name,
                        system_prompt=system_prompt,
                        user_content=user_content,
                        temperature=temperature,
                        extracted_data=extracted_data,
                    )

                st.session_state["report"] = report
                st.session_state["report_meta"] = {
                    "report_type": report_type,
                    "date_info": date_info,
                    "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
                }

            except Exception as e:
                st.error(f"❌ 生成报告过程中发生未预期错误：{e}")

    # ---------------- 报告展示与导出 ----------------
    if "report" in st.session_state:
        meta = st.session_state.get("report_meta", {})
        st.markdown("---")
        st.subheader(f"📄 {meta.get('report_type', '')}　|　{meta.get('date_info', '')}")

        st.markdown(st.session_state["report"])

        report_body = st.session_state["report"]
        md_doc = build_full_markdown(report_body, meta)
        html_doc = build_full_html(report_body, meta)

        st.markdown("**下载完整文档**")
        col_dl1, col_dl2, col_dl3 = st.columns(3)
        base_name = _safe_report_filename(meta, "").rstrip(".")

        with col_dl1:
            st.download_button(
                label="📄 Word 文档 (.docx)",
                data=build_docx_bytes(report_body, meta),
                file_name=f"{base_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with col_dl2:
            st.download_button(
                label="📝 Markdown (.md)",
                data=md_doc,
                file_name=f"{base_name}.md",
                mime="text/markdown",
                use_container_width=True,
            )
        with col_dl3:
            st.download_button(
                label="🌐 HTML 网页 (.html)",
                data=html_doc,
                file_name=f"{base_name}.html",
                mime="text/html",
                use_container_width=True,
            )

        st.markdown("**发送报告到邮箱**")
        default_recipients = _secret("email", "default_recipients")
        if not smtp_ready:
            st.warning(
                "请先在左侧「📧 邮件发信配置」中完成设置。"
                "Streamlit 云端若 SMTP 登录失败，请改用 **Resend API**。"
            )

        col_mail1, col_mail2 = st.columns([2, 1])
        with col_mail1:
            recipient_text = st.text_input(
                "收件人邮箱",
                value=default_recipients,
                placeholder="a@company.com, b@company.com",
                help="多个邮箱可用逗号、分号或换行分隔。",
            )
        with col_mail2:
            mail_format = st.selectbox(
                "附件格式",
                options=["docx", "html", "md"],
                format_func=lambda x: {"docx": "Word (.docx)", "html": "HTML (.html)", "md": "Markdown (.md)"}[x],
            )

        if st.button("📧 发送邮件", disabled=not smtp_ready, use_container_width=False):
            recipients = _parse_email_list(recipient_text)
            if not recipients:
                st.error("请填写至少一个有效的收件人邮箱。")
            else:
                with st.spinner("正在发送邮件..."):
                    ok, message = send_report_email(
                        recipients,
                        report_body,
                        meta,
                        attachment_format=mail_format,
                        smtp_override=smtp_override,
                        email_method=email_method,
                        resend_api_key=resend_api_key,
                        resend_from_addr=resend_from,
                    )
                if ok:
                    st.success(message)
                else:
                    st.error(message)


if __name__ == "__main__":
    main()
